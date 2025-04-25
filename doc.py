import docx
import re

def replace_placeholders(doc_path, output_path, replacements):
    """
    Replace placeholders in a Word document with actual values.
    Placeholders should be in format {{placeholder}}
    
    Args:
        doc_path: Path to the template document
        output_path: Path to save the filled document
        replacements: Dictionary with placeholder names and their values
    """
    doc = docx.Document(doc_path)
    
    # Regular expression to find placeholders like {{name}}
    pattern = r'\{\{([^}]+)\}\}'
    
    # Go through each paragraph in the document
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
            
        # Find all matches in this paragraph
        matches = re.findall(pattern, paragraph.text)
        
        # If we found placeholder(s), replace them
        if matches:
            text = paragraph.text
            
            # Make a copy of the runs to preserve formatting
            runs = [run._element for run in paragraph.runs]
            
            # Clear the paragraph
            for _ in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(p[0])
            
            # Replace all placeholders in the text
            for placeholder in matches:
                if placeholder in replacements:
                    replacement_value = replacements[placeholder]
                    text = text.replace('{{' + placeholder + '}}', replacement_value)
            
            # Add the modified text back to the paragraph
            paragraph.add_run(text)
    
    # Also search through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text:
                        continue
                        
                    # Find all matches in this paragraph
                    matches = re.findall(pattern, paragraph.text)
                    
                    # If we found placeholder(s), replace them
                    if matches:
                        text = paragraph.text
                        
                        # Make a copy of the runs to preserve formatting
                        runs = [run._element for run in paragraph.runs]
                        
                        # Clear the paragraph
                        for _ in range(len(paragraph.runs)):
                            p = paragraph._p
                            p.remove(p[0])
                        
                        # Replace all placeholders in the text
                        for placeholder in matches:
                            if placeholder in replacements:
                                replacement_value = replacements[placeholder]
                                text = text.replace('{{' + placeholder + '}}', replacement_value)
                        
                        # Add the modified text back to the paragraph
                        paragraph.add_run(text)
    
    # Save the document
    doc.save(output_path)
    return True

# Sample data to replace the placeholders
data = {
    'name': 'John Smith',
    'city': 'New York',
    'address': '123 Main Street',
    'date': 'April 25, 2025'
}

# Execute the replacement
replace_placeholders('2-.docx', 'filled_letter1.docx', data)

print("Document has been successfully created with filled placeholders!")